"""
OFAC NON-SDN Mapping Evidence Document Generator

Generates a Word document (.docx) showing field-by-field evidence with:
- XML snippet (raw source data)
- Mapping sheet reference (from OFAC_NON_SDN_Mapping.xlsx)
- MongoDB value (actual stored data in collection)
- Status (PASS / FAIL / MISSING)

For business stakeholders to review the data mapping accuracy.

Usage:
    python scripts/generate_ofac_evidence.py
    python scripts/generate_ofac_evidence.py --entity-id 15268
    python scripts/generate_ofac_evidence.py --sample 3

Requirements:
    pip install openpyxl pymongo python-docx
"""

import sys
import os
import json
import argparse
import random
from datetime import datetime
from collections import OrderedDict

try:
    import openpyxl
except ImportError:
    print("ERROR: pip install openpyxl"); sys.exit(1)

try:
    from pymongo import MongoClient
except ImportError:
    print("ERROR: pip install pymongo"); sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: pip install python-docx"); sys.exit(1)

import xml.etree.ElementTree as ET

# --- Config ---
DEFAULT_XML = os.path.join(os.environ.get("USERPROFILE", ""), "Downloads", "20260623T141041_cons_enhanced.xml")
DEFAULT_MAPPING = os.path.join(os.environ.get("USERPROFILE", ""), "Downloads", "OFAC_NON_SDN_Mapping.xlsx")
DEFAULT_OUTPUT = os.path.join(os.environ.get("USERPROFILE", ""), "Downloads", "OFAC_NON_SDN_Evidence_Document.docx")

MONGO_URI = "mongodb://qasaasuserrw:ZnTwAy0eTbaNdX1U@127.0.0.1:27017/?tls=true&directConnection=true&tlsInsecure=true"
MONGO_DATABASE = "screenDB"
MONGO_COLLECTION = "dataviumRegulatoryListHist"
LIST_ID = 94
NS = '{https://sanctionslistservice.ofac.treas.gov/api/PublicationPreview/exports/ENHANCED_XML}'


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color
    })
    shading_elm.append(shading)


def add_evidence_table(doc, rows_data, entity_id):
    """Add an evidence table for one entity to the document."""
    doc.add_heading(f'Entity ID: {entity_id}', level=2)

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header
    headers = ['#', 'Field (Mapping Sheet)', 'XML Source Path', 'XML Value', 'DB Field & Value', 'Status']
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(8)
        set_cell_shading(hdr_cells[i], '1A237E')
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for idx, row in enumerate(rows_data, 1):
        cells = table.add_row().cells
        cells[0].text = str(idx)
        cells[1].text = str(row.get('mapping_field', ''))[:60]
        cells[2].text = str(row.get('xml_path', ''))[:80]
        cells[3].text = str(row.get('xml_value', ''))[:100]
        cells[4].text = str(row.get('db_value', ''))[:100]
        status = row.get('status', 'PASS')
        cells[5].text = status

        # Color code status
        if status == 'PASS':
            set_cell_shading(cells[5], 'C8E6C9')
        elif status == 'FAIL':
            set_cell_shading(cells[5], 'FFCDD2')
        elif status == 'MISSING':
            set_cell_shading(cells[5], 'FFF3E0')

        # Font size
        for cell in cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(7.5)

    doc.add_paragraph('')  # spacing


def get_xml_snippet(entity_elem, path_hint):
    """Get a simplified XML snippet for evidence."""
    # Return a readable text representation of the relevant XML section
    lines = []

    def elem_to_str(elem, indent=0, max_depth=3):
        if indent > max_depth:
            return
        tag = elem.tag.replace(NS, '')
        text = (elem.text or '').strip()
        attribs = ' '.join(f'{k}="{v}"' for k, v in elem.attrib.items())
        if attribs:
            lines.append(f"{'  '*indent}<{tag} {attribs}>{text if text else ''}")
        else:
            lines.append(f"{'  '*indent}<{tag}>{text if text else ''}")
        for child in elem:
            elem_to_str(child, indent+1, max_depth)
        if not text and len(list(elem)) == 0:
            pass
        elif len(list(elem)) > 0:
            lines.append(f"{'  '*indent}</{tag}>")

    elem_to_str(entity_elem, max_depth=2)
    return '\n'.join(lines[:20])  # limit


def build_evidence_for_entity(entity_elem, db_doc, mappings):
    """Build field-by-field evidence data for one entity."""
    evidence = []
    eid = entity_elem.attrib.get('id', '')

    # Helper to add evidence row
    def add_row(mapping_field, xml_path, xml_value, db_field, db_value, status):
        evidence.append({
            'mapping_field': mapping_field,
            'xml_path': xml_path,
            'xml_value': xml_value if xml_value else '(empty)',
            'db_value': f"{db_field} = {db_value}" if db_value else f"{db_field} = (not found)",
            'status': status
        })

    # 1. Entity @id → sourceNaturalKey, listEntryId
    db_snk = db_doc.get('sourceNaturalKey', '') if db_doc else ''
    db_lei = db_doc.get('listEntryId', '') if db_doc else ''
    add_row('entity.@id → sourceNaturalKey', 'entity @id', eid,
            'sourceNaturalKey', db_snk, 'PASS' if str(eid) == str(db_snk) else 'FAIL')
    add_row('entity.@id → listEntryId', 'entity @id', eid,
            'listEntryId', db_lei, 'PASS' if str(eid) in str(db_lei) else 'FAIL')

    # 2. entityType
    gi = entity_elem.find(f'{NS}generalInfo')
    xml_type = ''
    if gi is not None:
        et = gi.find(f'{NS}entityType')
        if et is not None:
            xml_type = (et.text or '').strip()
    db_type = db_doc.get('entityTypeName', '') if db_doc else ''
    add_row('generalInfo.entityType → entityTypeName',
            'sanctionsData.entities.entity.generalInfo.entityType.#text', xml_type,
            'entityTypeName', db_type,
            'PASS' if xml_type.lower() == str(db_type).lower() else 'FAIL')

    # 3. Remarks → additionalInformation
    remarks_elem = gi.find(f'{NS}remarks') if gi is not None else None
    xml_remarks = (remarks_elem.text or '').strip() if remarks_elem is not None else ''
    db_remarks = db_doc.get('additionalInformation', '') if db_doc else ''
    if xml_remarks:
        add_row('generalInfo.remarks → additionalInformation',
                'entity.generalInfo.remarks', xml_remarks[:80],
                'additionalInformation', str(db_remarks)[:80],
                'PASS' if xml_remarks[:40] in str(db_remarks) else 'FAIL')

    # 4. Sanctions Lists
    sl_elem = entity_elem.find(f'{NS}sanctionsLists')
    db_sl = db_doc.get('sanctionListDetails', []) if db_doc else []
    if sl_elem is not None:
        for s in sl_elem.findall(f'{NS}sanctionsList'):
            xml_name = (s.text or '').strip()
            xml_date = s.attrib.get('datePublished', '')
            found = any(xml_name == d.get('sanctionsListName', '') for d in db_sl)
            db_match = next((d for d in db_sl if d.get('sanctionsListName') == xml_name), {})
            add_row('sanctionsList.#text → sanctionListDetails[].sanctionsListName',
                    f'sanctionsList @datePublished="{xml_date}"', xml_name,
                    'sanctionListDetails[].sanctionsListName', db_match.get('sanctionsListName', ''),
                    'PASS' if found else 'FAIL')
            if xml_date:
                add_row('sanctionsList.@datePublished → sanctionListDetails[].publishedDate',
                        'sanctionsList @datePublished', xml_date,
                        'sanctionListDetails[].publishedDate', db_match.get('publishedDate', ''),
                        'PASS' if xml_date == db_match.get('publishedDate', '') else 'FAIL')

    # 5. Sanctions Programs
    sp_elem = entity_elem.find(f'{NS}sanctionsPrograms')
    db_sp = db_doc.get('sanctionProgramDetailsList', []) if db_doc else []
    db_prog_names = [p.get('programName', '') for p in db_sp]
    if sp_elem is not None:
        for p in sp_elem.findall(f'{NS}sanctionsProgram'):
            xml_prog = (p.text or '').strip()
            add_row('sanctionsProgram.#text → sanctionProgramDetailsList[].programName',
                    'entity.sanctionsPrograms.sanctionsProgram', xml_prog,
                    'sanctionProgramDetailsList[].programName', xml_prog if xml_prog in db_prog_names else '(not found)',
                    'PASS' if xml_prog in db_prog_names else 'FAIL')

    # 6. Sanctions Types
    st_elem = entity_elem.find(f'{NS}sanctionsTypes')
    db_si = db_doc.get('sanctionImposedIndicatorsList', []) if db_doc else []
    if st_elem is not None:
        for t in st_elem.findall(f'{NS}sanctionsType'):
            xml_st = (t.text or '').strip()
            add_row('sanctionsType.#text → sanctionImposedIndicatorsList[]',
                    'entity.sanctionsTypes.sanctionsType', xml_st,
                    'sanctionImposedIndicatorsList', xml_st if xml_st in db_si else str(db_si),
                    'PASS' if xml_st in db_si else 'FAIL')

    # 7. Legal Authorities
    la_elem = entity_elem.find(f'{NS}legalAuthorities')
    db_la = db_doc.get('legalAuthority', []) if db_doc else []
    if not isinstance(db_la, list):
        db_la = [db_la] if db_la else []
    if la_elem is not None:
        for a in la_elem.findall(f'{NS}legalAuthority'):
            xml_la = (a.text or '').strip()
            found = xml_la in db_la or any(xml_la in str(d) or str(d) in xml_la for d in db_la)
            add_row('legalAuthority.#text → legalAuthority[]',
                    'entity.legalAuthorities.legalAuthority', xml_la[:60],
                    'legalAuthority', str(db_la)[:60],
                    'PASS' if found else 'FAIL')

    # 8. Names
    names_elem = entity_elem.find(f'{NS}names')
    db_names = db_doc.get('nameDetailsList', []) if db_doc else []
    if names_elem is not None:
        for name in names_elem.findall(f'{NS}name'):
            is_primary_elem = name.find(f'{NS}isPrimary')
            is_primary = (is_primary_elem.text or '').strip() if is_primary_elem is not None else ''
            is_primary = (is_primary or '').strip()
            alias_type_elem = name.find(f'{NS}aliasType')
            alias_type = (alias_type_elem.text or '').strip() if alias_type_elem is not None else ''
            is_low = name.find(f'{NS}isLowQuality')
            is_low_quality = (is_low.text or '').strip() if is_low is not None else 'false'

            translations = name.find(f'{NS}translations')
            if translations is not None:
                for tr in translations.findall(f'{NS}translation'):
                    xml_full = (tr.find(f'{NS}formattedFullName').text or '').strip() if tr.find(f'{NS}formattedFullName') is not None else ''
                    xml_first = (tr.find(f'{NS}formattedFirstName').text or '').strip() if tr.find(f'{NS}formattedFirstName') is not None else ''
                    xml_last = (tr.find(f'{NS}formattedLastName').text or '').strip() if tr.find(f'{NS}formattedLastName') is not None else ''
                    script_elem = tr.find(f'{NS}script')
                    xml_script = (script_elem.text or '').strip() if script_elem is not None else ''

                    # Find DB match
                    db_match = next((n for n in db_names if n.get('fullName', '') == xml_full), None)

                    add_row('names.translation.formattedFullName → nameDetailsList[].fullName',
                            f'name isPrimary={is_primary}', xml_full,
                            'nameDetailsList[].fullName', db_match.get('fullName', '(not found)') if db_match else '(not found)',
                            'PASS' if db_match else 'MISSING')

                    if xml_first:
                        add_row('names.translation.formattedFirstName → nameDetailsList[].firstName',
                                'formattedFirstName', xml_first,
                                'nameDetailsList[].firstName', db_match.get('firstName', '') if db_match else '',
                                'PASS' if db_match and xml_first == db_match.get('firstName', '') else 'FAIL' if db_match else 'MISSING')

                    if xml_last:
                        add_row('names.translation.formattedLastName → nameDetailsList[].lastName',
                                'formattedLastName', xml_last,
                                'nameDetailsList[].lastName', db_match.get('lastName', '') if db_match else '',
                                'PASS' if db_match and xml_last == db_match.get('lastName', '') else 'FAIL' if db_match else 'MISSING')

                    # nameType
                    expected_type = 'Primary' if is_primary == 'true' else alias_type
                    if db_match and expected_type:
                        db_nt = db_match.get('nameType', '')
                        match = expected_type.lower().replace('.', '') in db_nt.lower().replace('.', '') or db_nt.lower() in expected_type.lower()
                        add_row('name.isPrimary/aliasType → nameDetailsList[].nameType',
                                f'isPrimary={is_primary}, aliasType={alias_type}', expected_type,
                                'nameDetailsList[].nameType', db_nt,
                                'PASS' if match else 'FAIL')

                    # nameCategory
                    expected_cat = 'weak' if is_low_quality == 'true' else 'strong'
                    if db_match:
                        db_cat = db_match.get('nameCategory', '')
                        add_row('name.isLowQuality → nameDetailsList[].nameCategory',
                                f'isLowQuality={is_low_quality}', expected_cat,
                                'nameDetailsList[].nameCategory', db_cat,
                                'PASS' if expected_cat == db_cat else 'FAIL')

                    # script → originalScriptLanguage
                    if xml_script and db_match:
                        db_script = db_match.get('originalScriptLanguage', '')
                        add_row('translation.script → nameDetailsList[].originalScriptLanguage',
                                'script.#text', xml_script,
                                'nameDetailsList[].originalScriptLanguage', db_script,
                                'PASS' if xml_script == db_script else 'FAIL')

    # 9. Addresses
    addr_elem = entity_elem.find(f'{NS}addresses')
    db_addrs = db_doc.get('addressDetailsList', []) if db_doc else []
    if addr_elem is not None:
        for addr in addr_elem.findall(f'{NS}address'):
            country_elem = addr.find(f'{NS}country')
            xml_country = (country_elem.text or '').strip() if country_elem is not None else ''
            if xml_country:
                found = any(a.get('countryName', '') == xml_country for a in db_addrs)
                add_row('address.country → addressDetailsList[].countryName',
                        'address.country.#text', xml_country,
                        'addressDetailsList[].countryName', xml_country if found else '(not found)',
                        'PASS' if found else 'FAIL')

            translations = addr.find(f'{NS}translations')
            if translations is not None:
                for tr in translations.findall(f'{NS}translation'):
                    ap_elem = tr.find(f'{NS}addressParts')
                    if ap_elem is not None:
                        type_map = {'ADDRESS1': 'addressLine1', 'ADDRESS2': 'addressLine2',
                                    'ADDRESS3': 'addressLine3', 'CITY': 'city',
                                    'STATE/PROVINCE': 'stateOrProvince', 'POSTAL CODE': 'postalCode', 'REGION': 'region'}
                        for ap in ap_elem.findall(f'{NS}addressPart'):
                            ap_type_elem = ap.find(f'{NS}type')
                            ap_val_elem = ap.find(f'{NS}value')
                            if ap_type_elem is not None and ap_val_elem is not None:
                                ap_type = (ap_type_elem.text or '').strip()
                                ap_val = (ap_val_elem.text or '').strip()
                                db_field = type_map.get(ap_type, ap_type)
                                found = any(str(a.get(db_field, '')) == ap_val for a in db_addrs)
                                add_row(f'addressPart[{ap_type}] → addressDetailsList[].{db_field}',
                                        f'addressPart type="{ap_type}"', ap_val,
                                        f'addressDetailsList[].{db_field}', ap_val if found else '(not found)',
                                        'PASS' if found else 'FAIL')

    # 10. Features (Birthdate, Place of Birth, IDs, Citizenship, etc.)
    feat_elem = entity_elem.find(f'{NS}features')
    if feat_elem is not None:
        for feat in feat_elem.findall(f'{NS}feature'):
            ft_elem = feat.find(f'{NS}type')
            feat_type = (ft_elem.text or '').strip() if ft_elem is not None else ''
            val_elem = feat.find(f'{NS}value')
            feat_value = (val_elem.text or '').strip() if val_elem is not None else ''

            if feat_type == 'Birthdate' and feat_value:
                db_bd = db_doc.get('birthDateDetails', []) or db_doc.get('birthDateDetailsList', []) or []
                found = any(feat_value in str(b.get('date', '')) for b in db_bd)
                db_val = next((b.get('date', '') for b in db_bd if feat_value in str(b.get('date', ''))), '(not found)')
                add_row('feature[Birthdate].value → birthDateDetails[].date',
                        'feature type="Birthdate"', feat_value,
                        'birthDateDetails[].date', str(db_val),
                        'PASS' if found else 'MISSING')

            elif feat_type == 'Place of Birth' and feat_value:
                db_pob = db_doc.get('placeOfBirthDetails', []) or db_doc.get('placeOfBirthDetailsList', []) or []
                found = any(feat_value in str(p) for p in db_pob)
                add_row('feature[Place of Birth].value → placeOfBirthDetails[]',
                        'feature type="Place of Birth"', feat_value,
                        'placeOfBirthDetails', feat_value if found else '(not found)',
                        'PASS' if found else 'MISSING')

            elif 'Citizenship' in feat_type and feat_value:
                db_cit = db_doc.get('citizenshipDetails', []) or db_doc.get('citizenshipDetailsList', []) or []
                found = any(feat_value in str(c.get('countryName', '')) for c in db_cit)
                add_row(f'feature[{feat_type}].value → citizenshipDetails[].countryName',
                        f'feature type="{feat_type}"', feat_value,
                        'citizenshipDetails[].countryName', feat_value if found else '(not found)',
                        'PASS' if found else 'MISSING')

            elif feat_type and feat_value and feat_type not in ('Birthdate', 'Place of Birth') and 'Citizenship' not in feat_type:
                db_ids = db_doc.get('idNumberTypesList', []) if db_doc else []
                found = any(feat_value == str(i.get('idValue', '')) for i in db_ids)
                if found:
                    db_id = next((i for i in db_ids if feat_value == str(i.get('idValue', ''))), {})
                    add_row(f'feature[{feat_type}].value → idNumberTypesList[].idValue',
                            f'feature type="{feat_type}"', feat_value,
                            'idNumberTypesList[].idValue', db_id.get('idValue', ''),
                            'PASS')
                    # Also check idType
                    add_row(f'feature[{feat_type}].type → idNumberTypesList[].idType',
                            f'feature type="{feat_type}"', feat_type,
                            'idNumberTypesList[].idType', db_id.get('idType', ''),
                            'PASS' if feat_type == db_id.get('idType', '') or feat_type.lower() in db_id.get('idType', '').lower() else 'FAIL')
                else:
                    # Could be sourceSpecificInfo
                    db_ssi = db_doc.get('sourceSpecificInfoDetails', []) or db_doc.get('sourceSpecificInfoDetailsList', []) or []
                    ssi_found = any(feat_value in str(s.get('value', '')) for s in db_ssi)
                    if ssi_found:
                        add_row(f'feature[{feat_type}] → sourceSpecificInfoDetails[].value',
                                f'feature type="{feat_type}"', feat_value[:60],
                                'sourceSpecificInfoDetails[].value', feat_value[:60],
                                'PASS')
                    elif db_ids or db_ssi:
                        add_row(f'feature[{feat_type}].value → idNumberTypesList/sourceSpecificInfo',
                                f'feature type="{feat_type}"', feat_value[:60],
                                'DB lookup', '(not found)',
                                'MISSING')

    return evidence


def main():
    parser = argparse.ArgumentParser(description="Generate OFAC NON-SDN Evidence Document")
    parser.add_argument("--xml", default=DEFAULT_XML, help="Path to XML file")
    parser.add_argument("--mapping", default=DEFAULT_MAPPING, help="Path to mapping Excel")
    parser.add_argument("--output", "-o", default=DEFAULT_OUTPUT, help="Output .docx path")
    parser.add_argument("--sample", type=int, default=5, help="Number of entities (0=all)")
    parser.add_argument("--entity-id", type=str, help="Specific entity ID")
    args = parser.parse_args()

    if not os.path.exists(args.xml):
        print(f"ERROR: XML not found: {args.xml}"); sys.exit(1)
    if not os.path.exists(args.mapping):
        print(f"ERROR: Mapping not found: {args.mapping}"); sys.exit(1)

    # Load mapping
    print(f"Loading mapping: {args.mapping}")
    wb = openpyxl.load_workbook(args.mapping, read_only=True)
    ws = wb['OFAC NON SDN']
    mapping_rows = list(ws.iter_rows(values_only=True))
    print(f"  {len(mapping_rows)-1} mapping entries loaded")

    # Parse XML
    print(f"Parsing XML: {args.xml}")
    tree = ET.parse(args.xml)
    root = tree.getroot()
    entities_elem = root.find(f'{NS}entities')
    entity_list = {e.attrib.get('id'): e for e in entities_elem.findall(f'{NS}entity')}
    print(f"  {len(entity_list)} entities in XML")

    # Connect MongoDB
    print(f"Connecting to MongoDB...")
    client = MongoClient(MONGO_URI, serverSelectionTimeoutMS=10000)
    db = client[MONGO_DATABASE]
    collection = db[MONGO_COLLECTION]
    count = collection.count_documents({"listId": LIST_ID})
    print(f"  {count} documents in collection for listId={LIST_ID}")

    # Select entities
    if args.entity_id:
        entity_ids = [args.entity_id]
    elif args.sample > 0 and args.sample < len(entity_list):
        entity_ids = random.sample(list(entity_list.keys()), args.sample)
    else:
        entity_ids = list(entity_list.keys())

    print(f"\nGenerating evidence for {len(entity_ids)} entities...")

    # Create Word Document
    doc = Document()

    # Title
    title = doc.add_heading('OFAC NON-SDN Data Mapping Evidence Document', level=0)
    doc.add_paragraph(
        f'Collection: screenDB.{MONGO_COLLECTION} (listId={LIST_ID})\n'
        f'XML Source: {os.path.basename(args.xml)}\n'
        f'Mapping Sheet: {os.path.basename(args.mapping)} → "OFAC NON SDN" tab\n'
        f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}\n'
        f'Entities Verified: {len(entity_ids)} of {len(entity_list)}'
    )

    # Summary counts
    total_pass = 0
    total_fail = 0
    total_missing = 0

    for i, eid in enumerate(entity_ids, 1):
        entity_elem = entity_list.get(eid)
        if entity_elem is None:
            continue

        db_doc = collection.find_one({"listId": LIST_ID, "sourceNaturalKey": str(eid)})
        if db_doc is None:
            db_doc = collection.find_one({"listId": LIST_ID, "listEntryId": f"OFACNONSDN-{eid}"})

        evidence = build_evidence_for_entity(entity_elem, db_doc, mapping_rows)

        # Count statuses
        for row in evidence:
            if row['status'] == 'PASS':
                total_pass += 1
            elif row['status'] == 'FAIL':
                total_fail += 1
            else:
                total_missing += 1

        # Add page break between entities (except first)
        if i > 1:
            doc.add_page_break()

        add_evidence_table(doc, evidence, eid)

        status = "✅" if all(r['status'] == 'PASS' for r in evidence) else "❌"
        print(f"  [{i}/{len(entity_ids)}] Entity {eid}: {status} ({len(evidence)} fields)")

    # Add summary at the beginning (insert after title)
    doc.add_page_break()
    doc.add_heading('Verification Summary', level=1)
    total = total_pass + total_fail + total_missing
    summary_table = doc.add_table(rows=5, cols=2)
    summary_table.style = 'Table Grid'
    summary_data = [
        ('Total Field Checks', str(total)),
        ('✅ PASS', str(total_pass)),
        ('❌ FAIL', str(total_fail)),
        ('⚠️ MISSING', str(total_missing)),
        ('Pass Rate', f'{total_pass/total*100:.1f}%' if total > 0 else 'N/A'),
    ]
    for i, (label, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = value

    # Save
    doc.save(args.output)
    print(f"\n{'='*60}")
    print(f"Evidence document saved: {args.output}")
    print(f"  PASS: {total_pass} | FAIL: {total_fail} | MISSING: {total_missing}")
    print(f"  Pass Rate: {total_pass/total*100:.1f}%" if total > 0 else "")


if __name__ == "__main__":
    main()
